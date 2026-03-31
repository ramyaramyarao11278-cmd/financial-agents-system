from __future__ import annotations

import re
from pathlib import Path
from urllib.parse import unquote

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DOCUMENTS_DIR = PROJECT_ROOT / "documents"
FINAL_MD_PATH = DOCUMENTS_DIR / "完整论文_正式版.md"
OUTPUT_PATH = DOCUMENTS_DIR / "完整论文_正式版_组合排版.docx"
DEFAULT_TEMPLATE_DIR = Path(
    r"C:\Users\da983\Documents\xwechat_files\wxid_gc6r9mc9tebt22_e63d\temp\RWTemp\2026-03\35e1d01417cb3e03f875deece461ee41"
)


def resolve_template() -> Path:
    candidates = sorted(DEFAULT_TEMPLATE_DIR.glob("*.docx"))
    if not candidates:
        raise FileNotFoundError(f"No template docx found in {DEFAULT_TEMPLATE_DIR}")
    return candidates[0]


def extract_markdown_title(md_path: Path) -> str:
    content = md_path.read_text(encoding="utf-8")
    match = re.search(r"\*\*(.+?)\*\*", content)
    if not match:
        raise ValueError("Could not extract title from markdown")
    return match.group(1).strip()


def clean_inline_markdown(text: str) -> str:
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", text)
    text = text.replace("**", "")
    text = text.replace("`", "")
    return text.strip()


def normalize_windows_path(path_text: str) -> Path:
    normalized = unquote(path_text)
    if re.match(r"^/[A-Za-z]:/", normalized):
        normalized = normalized[1:]
    return Path(normalized)


def apply_run_font(run, *, size: float = 12.0, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:eastAsia"), "宋体")


def apply_cover_font(run) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)
    run.font.bold = True
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:eastAsia"), "黑体")


def format_body_paragraph(paragraph, *, first_line: bool = True, bold_prefix: str | None = None) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.first_line_indent = Pt(21) if first_line else Pt(0)

    text = paragraph.text
    paragraph.clear()
    if bold_prefix and text.startswith(bold_prefix):
        prefix_run = paragraph.add_run(bold_prefix)
        apply_run_font(prefix_run, bold=True)
        suffix = text[len(bold_prefix) :].lstrip()
        if suffix:
            suffix_run = paragraph.add_run(" " + suffix)
            apply_run_font(suffix_run)
        return

    run = paragraph.add_run(text)
    apply_run_font(run)


def format_reference_paragraph(paragraph) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.first_line_indent = Pt(0)
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(paragraph.text)
    paragraph.clear()
    run = paragraph.add_run(paragraph.text)
    apply_run_font(run)


def update_cover_title(doc: Document, title: str) -> None:
    non_empty = [p for p in doc.paragraphs if p.text.strip()]
    if len(non_empty) < 4:
        return

    title_paragraphs = non_empty[1:3]
    if "金融" in title:
        split_at = title.index("金融")
        lines = [title[:split_at], title[split_at:]]
    else:
        midpoint = max(1, len(title) // 2)
        lines = [title[:midpoint], title[midpoint:]]

    for paragraph, line in zip(title_paragraphs, lines):
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(line)
        apply_cover_font(run)


def insert_picture_with_caption(doc: Document, image_path: Path, caption: str) -> None:
    if not image_path.exists():
        warning = doc.add_paragraph(f"[缺图] {caption} -> {image_path}")
        format_body_paragraph(warning, first_line=False)
        return

    pic_para = doc.add_paragraph()
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_para.paragraph_format.first_line_indent = Pt(0)
    run = pic_para.add_run()
    run.add_picture(str(image_path), width=Cm(14.5))

    caption_para = doc.add_paragraph()
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_para.paragraph_format.first_line_indent = Pt(0)
    caption_run = caption_para.add_run(caption)
    apply_run_font(caption_run, size=10.5)


def extract_figure_info(line: str) -> tuple[Path, str] | None:
    path_match = re.search(r"\[[^\]]+\]\(([^)]+\.png)\)", line)
    title_match = re.search(r"图题可写为“([^”]+)”", line)
    if not path_match or not title_match:
        return None
    return normalize_windows_path(path_match.group(1)), title_match.group(1)


def clear_from_first_abstract(doc: Document) -> None:
    marker = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() in {"摘  要", "摘要"}:
            marker = paragraph
            break
    if marker is None:
        raise ValueError("Could not find abstract marker in template docx")

    body = doc._element.body
    remove_now = False
    for child in list(body.iterchildren()):
        if child == marker._element:
            remove_now = True
        if remove_now and child.tag != qn("w:sectPr"):
            body.remove(child)


def add_level_1_heading(doc: Document, text: str, add_page_break: bool) -> None:
    if add_page_break:
        doc.add_page_break()
    paragraph = doc.add_paragraph(style="Heading 1")
    paragraph.add_run(text)


def add_heading(doc: Document, level: int, text: str) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)


def add_toc_line(doc: Document, text: str, level: int) -> None:
    style_name = f"toc {min(level, 3)}"
    paragraph = doc.add_paragraph(style=style_name)
    run = paragraph.add_run(text)
    apply_run_font(run, size=12.0)


def generate_docx(template_path: Path, md_path: Path, output_path: Path) -> None:
    doc = Document(str(template_path))
    update_cover_title(doc, extract_markdown_title(md_path))
    clear_from_first_abstract(doc)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    start_idx = next(i for i, line in enumerate(lines) if line.strip() == "# 摘要")
    content_lines = lines[start_idx:]

    in_toc = False
    add_page_break_for_next_h1 = True

    for raw_line in content_lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            continue

        if stripped.startswith(">"):
            figure_info = extract_figure_info(stripped)
            if figure_info is not None:
                image_path, caption = figure_info
                insert_picture_with_caption(doc, image_path, caption)
            continue

        if stripped.startswith("# "):
            heading_text = stripped[2:].strip()
            display_text = {
                "摘要": "摘  要",
                "ABSTRACT": "Abstract",
                "目录": "目  录",
            }.get(heading_text, heading_text)
            add_level_1_heading(doc, display_text, add_page_break_for_next_h1)
            add_page_break_for_next_h1 = True
            in_toc = heading_text == "目录"
            continue

        if stripped.startswith("## "):
            add_heading(doc, 2, clean_inline_markdown(stripped[3:].strip()))
            in_toc = False
            continue

        if stripped.startswith("### "):
            add_heading(doc, 3, clean_inline_markdown(stripped[4:].strip()))
            in_toc = False
            continue

        bullet_match = re.match(r"^(\s*)-\s+(.*)$", line)
        if in_toc and bullet_match:
            indent_spaces = len(bullet_match.group(1))
            toc_level = indent_spaces // 2 + 1
            add_toc_line(doc, clean_inline_markdown(bullet_match.group(2)), toc_level)
            continue

        paragraph = doc.add_paragraph()
        paragraph.add_run(clean_inline_markdown(stripped))

        if stripped.startswith("**关键词：**") or stripped.startswith("关键词："):
            format_body_paragraph(paragraph, first_line=False, bold_prefix="关键词：")
        elif stripped.startswith("**Keywords:**") or stripped.startswith("Keywords:"):
            format_body_paragraph(paragraph, first_line=False, bold_prefix="Keywords:")
        elif re.match(r"^\d+\.\s", stripped):
            format_reference_paragraph(paragraph)
        else:
            format_body_paragraph(paragraph, first_line=True)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> None:
    template_path = resolve_template()
    generate_docx(template_path, FINAL_MD_PATH, OUTPUT_PATH)
    print(f"Generated: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
